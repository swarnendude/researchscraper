import re
import os
import uuid
import requests
from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import tempfile
from werkzeug.utils import secure_filename
import io
from datetime import datetime
import anthropic
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# PDF and DOCX parsing
import PyPDF2
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

# Initialize Anthropic client (reads ANTHROPIC_API_KEY from environment)
claude_client = anthropic.Anthropic()

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Output folder for saved markdown files
OUTPUT_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_stream):
    """Extract text and hyperlinks from PDF file."""
    try:
        reader = PyPDF2.PdfReader(file_stream)
        text = ''
        urls = set()

        for page in reader.pages:
            # Extract visible text
            text += page.extract_text() + '\n'

            # Extract hyperlinks from annotations
            if '/Annots' in page:
                annotations = page['/Annots']
                if annotations:
                    for annot in annotations:
                        annot_obj = annot.get_object()
                        if annot_obj.get('/Subtype') == '/Link':
                            if '/A' in annot_obj:
                                action = annot_obj['/A']
                                if '/URI' in action:
                                    uri = action['/URI']
                                    if uri and uri.startswith(('http://', 'https://')):
                                        urls.add(uri)

        # Append extracted URLs to text so they get picked up
        if urls:
            text += '\n\n--- Extracted Hyperlinks ---\n'
            for url in urls:
                text += url + '\n'

        return text
    except Exception as e:
        raise ValueError(f'Failed to parse PDF: {str(e)}')


def extract_text_from_docx(file_stream):
    """Extract text and hyperlinks from DOCX file."""
    try:
        # First, extract hyperlinks from the raw XML (more reliable)
        urls = set()
        file_stream.seek(0)

        with zipfile.ZipFile(file_stream, 'r') as zf:
            # Extract hyperlinks from document.xml.rels
            if 'word/_rels/document.xml.rels' in zf.namelist():
                rels_xml = zf.read('word/_rels/document.xml.rels')
                rels_root = ET.fromstring(rels_xml)
                for rel in rels_root.iter():
                    if rel.get('Type') and 'hyperlink' in rel.get('Type', '').lower():
                        target = rel.get('Target', '')
                        if target.startswith(('http://', 'https://')):
                            urls.add(target)

            # Also check for URLs in the document text itself
            if 'word/document.xml' in zf.namelist():
                doc_xml = zf.read('word/document.xml')
                # Find any http/https URLs in the raw XML
                url_matches = re.findall(r'https?://[^\s<>"\']+', doc_xml.decode('utf-8', errors='ignore'))
                for url in url_matches:
                    url = url.rstrip('.,;:!?)')
                    if url.startswith(('http://', 'https://')):
                        urls.add(url)

        # Now extract text using python-docx
        file_stream.seek(0)
        doc = Document(file_stream)
        text = ''

        for para in doc.paragraphs:
            text += para.text + '\n'

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + '\n'

        # Append extracted URLs to text so they get picked up
        if urls:
            text += '\n\n--- Extracted Hyperlinks ---\n'
            for url in urls:
                text += url + '\n'

        return text
    except Exception as e:
        raise ValueError(f'Failed to parse DOCX: {str(e)}')

# Store scraping results temporarily
scrape_results = {}


def extract_urls(text):
    """Extract URLs from text content."""
    url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
    urls = re.findall(url_pattern, text)
    # Clean up URLs (remove trailing punctuation)
    cleaned = []
    for url in urls:
        url = url.rstrip('.,;:!?)')
        if url and urlparse(url).netloc:
            cleaned.append(url)
    return list(dict.fromkeys(cleaned))  # Remove duplicates, preserve order


def generate_title_with_claude(content):
    """Use Claude to generate a short, descriptive title based on content."""
    try:
        # Take first 3000 chars of content for context
        content_preview = content[:3000] if len(content) > 3000 else content

        message = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=100,
            messages=[
                {
                    "role": "user",
                    "content": f"""Based on the following scraped content, generate a short, descriptive title (3-6 words) that captures the main topic. Return ONLY the title, nothing else.

Content:
{content_preview}"""
                }
            ]
        )

        title = message.content[0].text.strip()
        # Clean up the title - remove quotes if present
        title = title.strip('"\'')
        # Limit length
        if len(title) > 60:
            title = title[:60]
        return title
    except Exception as e:
        print(f"Error generating title with Claude: {e}")
        return None


def generate_filename_from_title(title):
    """Convert a title to a safe filename."""
    if not title:
        return 'scraped_content.md'

    # Remove special characters, keep alphanumeric, spaces, hyphens
    clean_title = re.sub(r'[^\w\s\-]', '', title)

    # Replace multiple spaces with single space
    clean_title = re.sub(r'\s+', ' ', clean_title).strip()

    # Truncate if too long (max 50 chars for filename)
    if len(clean_title) > 50:
        clean_title = clean_title[:50].rsplit(' ', 1)[0]

    # Replace spaces with underscores for filename
    clean_title = clean_title.replace(' ', '_')

    if not clean_title:
        return 'scraped_content.md'

    return f'{clean_title}.md'


def generate_filename_from_results(results):
    """Generate a meaningful filename based on scraped content titles."""
    # Collect successful titles
    titles = [r['title'] for r in results if r and r.get('success') and r.get('title')]

    if not titles:
        return 'scraped_content.md'

    # Use the first successful title as the base
    title = titles[0]

    # Clean the title for use as filename
    # Remove common suffixes like "| Site Name", "- Company", etc.
    title = re.split(r'\s*[\|\-\u2013\u2014]\s*', title)[0].strip()

    # Remove special characters, keep alphanumeric, spaces, hyphens
    title = re.sub(r'[^\w\s\-]', '', title)

    # Replace multiple spaces with single space
    title = re.sub(r'\s+', ' ', title).strip()

    # Truncate if too long (max 50 chars for filename)
    if len(title) > 50:
        title = title[:50].rsplit(' ', 1)[0]  # Cut at word boundary

    # Replace spaces with underscores for filename
    title = title.replace(' ', '_')

    # Fallback if title became empty
    if not title:
        return 'scraped_content.md'

    return f'{title}.md'


def scrape_url(url, timeout=10):
    """Scrape content from a single URL."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=timeout)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')

        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()

        # Get title
        title = soup.title.string if soup.title else url

        # Get main content - try common content containers
        content = None
        for selector in ['article', 'main', '.content', '#content', '.post', '.article']:
            content = soup.select_one(selector)
            if content:
                break

        if not content:
            content = soup.body if soup.body else soup

        # Extract text
        text = content.get_text(separator='\n', strip=True)

        # Clean up excessive whitespace
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n\n'.join(lines)

        # Limit content length
        if len(text) > 5000:
            text = text[:5000] + '\n\n[Content truncated...]'

        return {
            'success': True,
            'title': title.strip() if title else url,
            'content': text,
            'url': url
        }
    except Exception as e:
        return {
            'success': False,
            'title': url,
            'content': f'Error scraping: {str(e)}',
            'url': url
        }


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and extract text."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type. Only PDF and DOCX allowed'}), 400

    try:
        filename = secure_filename(file.filename)
        ext = filename.rsplit('.', 1)[1].lower()

        # Read file into memory
        file_stream = io.BytesIO(file.read())

        if ext == 'pdf':
            text = extract_text_from_pdf(file_stream)
        elif ext == 'docx':
            text = extract_text_from_docx(file_stream)
        else:
            return jsonify({'error': 'Unsupported file type'}), 400

        return jsonify({'text': text, 'filename': filename})

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': f'Failed to process file: {str(e)}'}), 500


@app.route('/scrape', methods=['POST'])
def scrape():
    data = request.json
    text = data.get('text', '')
    filename = data.get('filename', '')

    urls = extract_urls(text)
    if not urls:
        return jsonify({'error': 'No URLs found in the text'}), 400

    # Generate session ID
    session_id = str(uuid.uuid4())
    scrape_results[session_id] = {
        'urls': urls,
        'results': [],
        'total': len(urls),
        'completed': 0,
        'filename': filename
    }

    return jsonify({
        'session_id': session_id,
        'total': len(urls),
        'urls': urls
    })


@app.route('/scrape/<session_id>/<int:index>', methods=['POST'])
def scrape_single(session_id, index):
    if session_id not in scrape_results:
        return jsonify({'error': 'Session not found'}), 404

    session = scrape_results[session_id]
    if index >= len(session['urls']):
        return jsonify({'error': 'Invalid index'}), 400

    url = session['urls'][index]
    result = scrape_url(url)

    # Store result
    while len(session['results']) <= index:
        session['results'].append(None)
    session['results'][index] = result
    session['completed'] = sum(1 for r in session['results'] if r is not None)

    return jsonify({
        'result': result,
        'completed': session['completed'],
        'total': session['total']
    })


def get_unique_filename(folder, filename):
    """Generate a unique filename by adding a number suffix if file exists."""
    base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
    extension = '.md'

    # Check if file exists, if so add number suffix
    final_filename = filename
    counter = 1
    while os.path.exists(os.path.join(folder, final_filename)):
        final_filename = f'{base_name}_{counter}{extension}'
        counter += 1

    return final_filename


@app.route('/download/<session_id>')
def download(session_id):
    if session_id not in scrape_results:
        return jsonify({'error': 'Session not found'}), 404

    session = scrape_results[session_id]

    # Collect all content for title generation
    all_content = ""
    for result in session['results']:
        if result and result.get('success'):
            all_content += result.get('content', '') + "\n\n"

    # Generate title using Claude
    generated_title = generate_title_with_claude(all_content)
    if not generated_title:
        generated_title = "Scraped Content"

    # Build markdown document with generated title
    markdown = f'# {generated_title}\n\n'
    markdown += f'Total URLs: {session["total"]}\n\n---\n\n'

    for result in session['results']:
        if result:
            status = '✓' if result['success'] else '✗'
            markdown += f'## {status} {result["title"]}\n\n'
            markdown += f'**URL:** {result["url"]}\n\n'
            markdown += result['content']
            markdown += '\n\n---\n\n'

    # Use uploaded filename if available, otherwise generate from Claude title
    original_filename = session.get('filename', '')
    if original_filename:
        # Remove extension and add .md
        base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
        download_filename = f'Scraped_{base_name}.md'
    else:
        # Generate filename from Claude-generated title
        download_filename = f'Scraped_{generate_filename_from_title(generated_title)}'
        # Ensure .md extension
        if not download_filename.endswith('.md'):
            download_filename += '.md'

    # Ensure unique filename in output folder
    download_filename = get_unique_filename(OUTPUT_FOLDER, download_filename)

    # Save to output folder
    output_path = os.path.join(OUTPUT_FOLDER, download_filename)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown)

    return send_file(
        output_path,
        as_attachment=True,
        download_name=download_filename,
        mimetype='text/markdown'
    )


@app.route('/history')
def get_history():
    """Get list of all generated markdown files."""
    files = []
    for filename in os.listdir(OUTPUT_FOLDER):
        if filename.endswith('.md'):
            filepath = os.path.join(OUTPUT_FOLDER, filename)
            stat = os.stat(filepath)

            # Default title from filename
            title = filename.replace('.md', '').replace('_', ' ').replace('Scraped ', '')

            # Try to extract the Claude-generated title from the first line (# Title)
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    first_line = f.readline().strip()
                    if first_line.startswith('# '):
                        title = first_line[2:].strip()
                        if len(title) > 60:
                            title = title[:60] + '...'
            except:
                pass

            files.append({
                'filename': filename,
                'title': title,
                'created': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'size': stat.st_size
            })

    # Sort by creation time, newest first
    files.sort(key=lambda x: x['created'], reverse=True)
    return jsonify(files)


@app.route('/history/download/<path:filename>')
def download_history_file(filename):
    """Download a specific file from history."""
    # Only allow .md files and prevent directory traversal
    if '..' in filename or filename.startswith('/'):
        return jsonify({'error': 'Invalid filename'}), 400

    if not filename.endswith('.md'):
        return jsonify({'error': 'Invalid file type'}), 400

    return send_from_directory(
        OUTPUT_FOLDER,
        filename,
        as_attachment=True,
        mimetype='text/markdown'
    )


@app.route('/history/view/<path:filename>')
def view_history_file(filename):
    """View a specific file from history in browser."""
    # Only allow .md files and prevent directory traversal
    if '..' in filename or filename.startswith('/'):
        return jsonify({'error': 'Invalid filename'}), 400

    if not filename.endswith('.md'):
        return jsonify({'error': 'Invalid file type'}), 400

    return send_from_directory(
        OUTPUT_FOLDER,
        filename,
        as_attachment=False,
        mimetype='text/plain; charset=utf-8'
    )


@app.route('/history/delete/<path:filename>', methods=['POST'])
def delete_history_file(filename):
    """Delete a specific file from history."""
    # Only allow .md files and prevent directory traversal
    if '..' in filename or filename.startswith('/'):
        return jsonify({'error': 'Invalid filename'}), 400

    if not filename.endswith('.md'):
        return jsonify({'error': 'Invalid file type'}), 400

    filepath = os.path.join(OUTPUT_FOLDER, filename)

    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404

    try:
        os.remove(filepath)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True)
